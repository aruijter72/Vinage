#!/usr/bin/env python3
"""Genereert 'Partner-config handleiding.docx' uit deze beschrijving.
Hergenereren:  python3 docs/generate-partner-config-docx.py
Vereist:       pip install python-docx
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = os.path.join(os.path.dirname(__file__), "Partner-config handleiding.docx")
BURGUNDY = RGBColor(0x7A, 0x25, 0x35)


def shade(cell, hexcolor):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        shade(c, "7A2535")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(val))
            r.font.size = Pt(10)
    return t


def mono(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(11)
    return p


doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

title = doc.add_heading("Vinage — Partner-importers aan/uit zetten", level=0)
for run in title.runs:
    run.font.color.rgb = BURGUNDY

doc.add_paragraph(
    "Wijn & Spijs toont suggesties van partner-importers (Okhuysen, SooD, …). "
    "Welke partners zichtbaar zijn, kun je live aan- of uitzetten via de "
    "Firebase Console — zonder code-wijziging of nieuwe deploy."
)
p = doc.add_paragraph()
p.add_run("Waar: ").bold = True
p.add_run("Firebase Console → project ")
p.add_run("vinage-85fd8").bold = True
p.add_run(" → Firestore Database.")

doc.add_heading("1. Het document", level=1)
table(doc, ["Onderdeel", "Waarde"], [
    ["Collectie", "config"],
    ["Document-ID", "importers"],
    ["Pad", "config/importers"],
])

doc.add_heading("2. Eenmalig aanmaken", level=1)
for i, step in enumerate([
    'Klik “Start collection” / “+ Add collection” → Collection ID: config → Next.',
    "Document ID: importers  (exact, kleine letters).",
    "Voeg per partner een veld toe (zie tabel hieronder).",
    "Klik Save.",
], 1):
    doc.add_paragraph(f"{i}. {step}")

table(doc, ["Field name", "Type", "Value"], [
    ["sood", "boolean", "true  of  false"],
    ["okhuysen", "boolean", "true  of  false"],
])
note = doc.add_paragraph()
note.add_run("Let op: ").bold = True
note.add_run('het veldtype moet boolean zijn (de true/false-schakelaar), '
             'niet de tekst "true".')

doc.add_heading("3. Daarna wijzigen", level=1)
doc.add_paragraph(
    "Open config/importers, klik op het veld, wissel de boolean, klik Update. "
    "De wijziging is vrijwel direct live (de app luistert via een Firestore "
    "snapshot); uiterlijk zichtbaar bij de volgende Wijn & Spijs-zoekopdracht "
    "of na herladen. Geen deploy nodig."
)

doc.add_heading("Wat de waarden doen", level=1)
table(doc, ["Situatie", "Gedrag"], [
    ["Veld = true", "Partner wordt getoond"],
    ["Veld = false", "Partner wordt verborgen in Wijn & Spijs"],
    ["Veld/document bestaat niet", "Fallback naar standaard in de code (= aan)"],
    ["Offline of niet ingelogd", "Zelfde fallback: standaard uit de code (= aan)"],
])

doc.add_heading("Snelle referentie", level=1)
table(doc, ["Wil je…", "Zet veld", "Op waarde"], [
    ["SooD verbergen", "sood", "false"],
    ["SooD tonen", "sood", "true (of veld weghalen)"],
    ["Okhuysen verbergen", "okhuysen", "false"],
    ["Okhuysen tonen", "okhuysen", "true (of veld weghalen)"],
    ["Alles standaard (aan)", "—", "document niet aanmaken"],
])

doc.add_heading("Voorbeeld — SooD tijdelijk verbergen", level=1)
doc.add_paragraph(
    "Zet in config/importers het veld sood op false. Weer tonen: op true "
    "zetten of het veld verwijderen (dan geldt de code-standaard = aan)."
)
doc.add_paragraph("Document-inhoud (referentie):")
mono(doc, '{ "okhuysen": true, "sood": false }')
ref = doc.add_paragraph()
ref.add_run("De Firestore Console heeft geen “JSON plakken”-knop — voer de "
            "velden los in. Bovenstaand JSON-blok is alleen ter referentie / "
            "voor tooling (Admin SDK of importscript).").italic = True

doc.add_heading("Nieuwe partner later toevoegen", level=1)
doc.add_paragraph(
    "Komt er een derde partner bij, dan krijgt die een eigen id in "
    "js/importers.js (kleine letters). Voeg in hetzelfde document "
    "config/importers een veld toe met dat id als naam, type boolean."
)

doc.add_heading("Belangrijke aandachtspunten", level=1)
for b in [
    "Veldtype = boolean (niet de string \"true\").",
    "Document-ID en veldnamen exact en in kleine letters: importers, sood, "
    "okhuysen (= de partner-id uit js/importers.js).",
    "Schrijven kan alleen via de Console (firestore.rules: config is publiek "
    "leesbaar, schrijven geblokkeerd). Inhoud is niet gevoelig.",
    "Zolang het document niet bestaat, blijft alles zoals in de code (alle "
    "partners zichtbaar). Alleen aanmaken wanneer je iets wilt uitschakelen.",
]:
    doc.add_paragraph(b, style="List Bullet")

doc.add_heading("Technische verwijzingen (onderhoud)", level=1)
table(doc, ["Onderdeel", "Locatie"], [
    ["Partnerdefinities + catalogus + logo", "js/importers.js (IMPORTERS)"],
    ["Config inlezen + fallback", "js/sync.js (importerConfig, isImporterActive)"],
    ["Toepassing in Wijn & Spijs", "js/app.js (_buildImporterSuggestions)"],
    ["Beveiliging", "firestore.rules (match /config/{docId})"],
])

doc.save(OUT)
print("Geschreven:", OUT)
